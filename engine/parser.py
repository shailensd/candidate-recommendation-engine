import pdfplumber
from docx import Document
import re
import logging

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file):
    """
    Extracts and concatenates text from all pages of a PDF file.
    Returns empty string on failure.
    """
    try:
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        logger.error(f"Failed to read PDF with pdfplumber: {e}")
        return ""

def extract_text_from_docx(file):
    """
    Extracts text from both paragraphs and tables in a DOCX file.
    Maintains document structure by joining with newlines.
    """
    try:
        doc = Document(file)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_parts.append(text)
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to read DOCX: {e}")
        return ""

def clean_text(text):
    """
    Normalizes text by:
    1. Collapsing multiple whitespace into single space
    2. Removing special chars except basic punctuation
    3. Trimming leading/trailing whitespace
    """
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s\.\,\!\?\-\+\=\&\|\:\;\(\)\[\]\{\}]', '', text)
    return text.strip()

def extract_candidate_info(text):
    """
    Parses resume text to extract candidate details:
    - Name: First non-header line that looks like a name
    - Email: First matching email address
    - Phone: First matching phone number in various formats
    Returns fallback values if information not found.
    """
    if not text:
        return {
            "name": "Unknown Candidate",
            "email": "No email found",
            "phone": "No phone found",
            "full_text": ""
        }

    text = text.strip()
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # Extract name from first line that looks like a name
    name = "Unknown Candidate"
    for line in lines:
        if line and not line.startswith(('http', 'www', '@')) and len(line) < 100:
            # Skip section headers and contact info
            if any(skip in line.lower() for skip in ['phone', 'email', 'address', 'experience', 'education', 'skills']):
                continue
            # Match name pattern: letters, spaces, hyphens, apostrophes, 1-4 words
            if re.match(r'^[A-Za-z\s\-\.\']+$', line) and len(line.split()) <= 4:
                if line.isupper():
                    words = line.split()
                    if 2 <= len(words) <= 4:  # Multi-word all-caps likely name
                        name = line
                        break
                    elif len(words) == 1:  # Single word likely section header
                        continue
                else:
                    name = line
                    break

    # Extract email using standard email regex
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    email = email_match.group() if email_match else "No email found"

    # Extract phone using comprehensive format patterns
    phone_patterns = [
        # North American formats
        r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}(?:\s*(?:x|ext\.?|extension)\s*\d+)?',  # (408) 627-2229, with optional extension
        r'\+?1[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # +1 (408) 627-2229
        r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # 408-627-2229 or 408.627.2229
        
        # International formats
        r'\+44\s?\d{2,5}\s?\d{6}',  # UK: +44 7911 123456
        r'\+61\s?\d\s?\d{4}\s?\d{4}',  # Australia: +61 4 1234 5678
        r'\+86\s?\d{2,3}\s?\d{4}\s?\d{4}',  # China: +86 123 4567 8901
        
        # European formats
        r'\+\d{2}[-.\s]?\d{1,2}[-.\s]?\d{2,3}[-.\s]?\d{2}[-.\s]?\d{2}',  # +33 1 23 45 67 89
        r'\+\d{2}[-.\s]?\d{9,10}',  # +49 1234567890
        
        # Asian formats
        r'\+\d{2}[-.\s]?\d{3,5}[-.\s]?\d{4}',  # Japanese/Korean
        r'\+\d{2}[-.\s]?\d{5}[-.\s]?\d{5}',  # Indian: +91 98765 43210
        
        # Generic formats
        r'\+?\d{1}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # Generic international
        r'\d{10}',  # Plain 10 digits
        
        # With extensions
        r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}(?:\s*(?:x|ext\.?|extension)\s*\d+)?'  # Any format with extension
    ]
    
    phone = "No phone found"
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            phone = phone_match.group()
            break

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "full_text": text
    }